"""Unit tests for export_service.py — conversion core, markers, detection."""

import io
import re

import pytest

import export_service as es


# ---------------------------------------------------------------------------
# strip_llm_markup
# ---------------------------------------------------------------------------

def test_strip_thought_blocks():
    src = "До <think>секретные мысли</think>после"
    assert es.strip_llm_markup(src) == "До после"


def test_strip_unclosed_think_and_status():
    src = "Ответ готов.\n[STATUS: searching]\n<think>обрыв стрима"
    out = es.strip_llm_markup(src)
    assert out == "Ответ готов."


def test_strip_file_markers():
    src = 'Смотри файл.\n\n<<<FILE name="a.pdf" kind="pdf">>>\nсодержимое\n<<<END FILE>>>'
    assert es.strip_llm_markup(src) == "Смотри файл."


def test_strip_export_markers():
    src = 'Готово.\n\n<<<EXPORT format="pdf" message_id="5" name="x.pdf">>>'
    assert es.strip_llm_markup(src) == "Готово."


def test_neutralizes_active_html():
    src = "текст <script>alert(1)</script> <iframe src='http://evil'> конец"
    out = es.strip_llm_markup(src)
    assert "<script" not in out
    assert "<iframe" not in out
    assert "&lt;iframe" in out


# ---------------------------------------------------------------------------
# markdown → HTML
# ---------------------------------------------------------------------------

TABLE_MD = (
    "# Заголовок\n\n"
    "| Направление | Что включено |\n"
    "|---|---|\n"
    "| **Анализ** | Оценка бизнес-модели |\n"
    "| Рынок | Объём, тренды |\n"
)


def test_markdown_table_to_html():
    html = es.markdown_to_html(TABLE_MD)
    assert "<table>" in html
    assert "<th>Направление</th>" in html
    assert "<strong>Анализ</strong>" in html


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def test_docx_contains_real_table_and_heading():
    from docx import Document

    html = es.markdown_to_html(TABLE_MD)
    data = es.html_to_docx(html, title="Тест", subtitle="дата")
    doc = Document(io.BytesIO(data))
    assert len(doc.tables) == 1
    table = doc.tables[0]
    assert table.cell(0, 0).text.strip() == "Направление"
    assert table.cell(1, 0).text.strip() == "Анализ"
    texts = [p.text for p in doc.paragraphs]
    assert "Тест" in texts
    assert any("Заголовок" in t for t in texts)


def test_docx_lists_and_sources():
    from docx import Document

    html = es.markdown_to_html("- один\n- два\n  - вложенный\n\n1. раз\n2. два\n")
    data = es.html_to_docx(html, sources=[{"title": "Пример", "url": "https://example.com"}])
    doc = Document(io.BytesIO(data))
    texts = [p.text for p in doc.paragraphs]
    assert any("один" in t for t in texts)
    assert any("вложенный" in t for t in texts)
    assert any("Источники" in t for t in texts)


# ---------------------------------------------------------------------------
# TXT
# ---------------------------------------------------------------------------

def test_txt_ascii_table_grid_alignment():
    html = es.markdown_to_html(TABLE_MD)
    txt = es.html_to_txt(html)
    lines = txt.splitlines()
    border_lines = [ln for ln in lines if ln.startswith("+")]
    pipe_lines = [ln for ln in lines if ln.startswith("|")]
    assert border_lines and pipe_lines
    # all grid rows share one width
    widths = {len(ln) for ln in border_lines + pipe_lines}
    assert len(widths) == 1
    # header separated with '='
    assert any(set(ln) <= {"+", "="} for ln in border_lines)
    assert "Направление" in txt


def test_txt_long_cell_wraps():
    long_cell = "очень " * 30
    md = f"| A | B |\n|---|---|\n| {long_cell} | x |\n"
    txt = es.html_to_txt(es.markdown_to_html(md))
    for ln in txt.splitlines():
        if ln.startswith(("|", "+")):
            assert len(ln) <= es._TXT_CELL_MAX * 2 + 10


def test_txt_headings_and_lists():
    txt = es.html_to_txt(es.markdown_to_html("# Тема\n\nАбзац.\n\n- пункт один\n- пункт два\n"))
    assert "ТЕМА" in txt
    assert "- пункт один" in txt


# ---------------------------------------------------------------------------
# Filenames / Content-Disposition
# ---------------------------------------------------------------------------

def test_suggest_filename_cleans_markdown():
    name = es.suggest_filename("## Юнит-экономика *стартапа*", "pdf")
    assert name.endswith(".pdf")
    assert "Юнит-экономика стартапа" in name
    assert "#" not in name and "*" not in name


def test_content_disposition_rfc5987():
    cd = es.build_content_disposition("Ответ Pitchy — тест.pdf")
    assert cd.startswith("attachment; ")
    m = re.search(r'filename="([^"]+)"', cd)
    assert m and m.group(1).isascii()
    assert "filename*=UTF-8''" in cd
    assert "%20" in cd or "%E2" in cd  # urlencoded utf-8


def test_transliterate():
    assert es.transliterate("Юнит-экономика") == "Yunit-ekonomika"


# ---------------------------------------------------------------------------
# Markers
# ---------------------------------------------------------------------------

def test_export_marker_roundtrip():
    marker = es.build_export_marker("pdf", 'отчёт "финал".pdf', message_id=42)
    m = es.EXPORT_MARKER_RE.search(marker)
    assert m
    assert m.group(1) == "pdf"
    assert m.group(2) == "42"
    assert '"' not in m.group(3)


def test_export_marker_without_message_id():
    marker = es.build_export_marker("docx", "файл.docx")
    m = es.EXPORT_MARKER_RE.search(marker)
    assert m and m.group(2) is None and m.group(1) == "docx"


# ---------------------------------------------------------------------------
# detect_export_request (regex prefilter)
# ---------------------------------------------------------------------------

@pytest.mark.parametrize(
    "text",
    [
        "сохрани прошлый ответ в pdf",
        "Собери следующий ответ в txt файле",
        "экспортируй это в ворд",
        "сделай мне маркдаун файл из ответа",
        "скачать ответ как docx",
        "save this answer as markdown",
    ],
)
def test_detect_export_positive(text):
    assert es.detect_export_request(text) is True


@pytest.mark.parametrize(
    "text",
    [
        "расскажи про юнит-экономику",
        "какие есть гранты для стартапов?",
        "что такое pdf-формат?",  # формат без глагола действия
        "",
    ],
)
def test_detect_export_negative(text):
    assert es.detect_export_request(text) is False


# ---------------------------------------------------------------------------
# Facade
# ---------------------------------------------------------------------------

def test_render_md_passthrough_with_sources():
    data, mime, ext = es.render_message_export(
        TABLE_MD, "md", title="Тест", sources=[{"title": "Src", "url": "https://e.com"}]
    )
    text = data.decode("utf-8")
    assert "| Направление |" in text
    assert "## Источники" in text
    assert mime.startswith("text/markdown")
    assert ext == "md"


def test_render_txt_and_docx():
    for fmt in ("txt", "docx"):
        data, mime, ext = es.render_message_export(TABLE_MD, fmt, title="Тест")
        assert len(data) > 100
        assert ext == fmt


def test_render_unknown_format_raises():
    with pytest.raises(ValueError):
        es.render_message_export("x", "exe")


def test_render_pdf_unavailable_raises_cleanly():
    if es.pdf_available():
        data, mime, _ = es.render_message_export(TABLE_MD, "pdf", title="Тест")
        assert data[:5] == b"%PDF-"
        assert mime == "application/pdf"
    else:
        with pytest.raises(es.ExportUnavailable):
            es.render_message_export(TABLE_MD, "pdf", title="Тест")
