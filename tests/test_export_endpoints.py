"""Endpoint tests for GET /chat/messages/{id}/export."""

import importlib
import io
import sys
import uuid
from unittest import mock

import pytest
from fastapi.testclient import TestClient

# Mock rag module to avoid chromadb import issues (same as test_auth.py)
sys.modules["rag"] = mock.MagicMock()

TABLE_ANSWER = (
    "## Чем могу помочь\n\n"
    "| Направление | Что включено |\n"
    "|---|---|\n"
    "| Анализ стартапа | Оценка бизнес-модели |\n"
    "| Рыночный анализ | Объём рынка, тренды |\n"
)


def _register_and_login(client) -> str:
    email = f"exp_{uuid.uuid4()}@example.com"
    # Код подтверждения генерирует routers/auth.py — мокаем его random,
    # чтобы код был предсказуемым ("555555").
    with mock.patch("routers.auth.random.randint", return_value=5):
        res = client.post(
            "/auth/register",
            json={
                "email": email,
                "name": "Exporter",
                "password": "Pass1234",
                "accept_privacy": True,
                "accept_cookies": True,
            },
        )
        assert res.status_code == 200
        res = client.post("/auth/verify-email", json={"email": email, "code": "555555"})
        assert res.status_code == 200
    return res.json()["access_token"]


def _auth_headers(token: str) -> dict:
    """Auth идёт ТОЛЬКО по httpOnly-cookie (Bearer-фолбэк убран, см. auth.py).
    Кука логина ставится с domain=.pitchy.pro и не сохраняется в jar
    TestClient — поэтому шлём её явным заголовком."""
    return {"Cookie": f"access_token={token}"}


def _make_assistant_message(token: str, client, content: str = TABLE_ANSWER) -> int:
    """Creates a chat session via API, then inserts an assistant reply
    directly into the DB (the LLM pipeline is not under test here)."""
    res = client.post(
        "/chat/sessions",
        json={"title": "Юнит-экономика"},
        headers=_auth_headers(token),
    )
    assert res.status_code == 200, res.text
    session_id = res.json()["id"]

    from db import SessionLocal
    from models import ChatMessage

    db = SessionLocal()
    try:
        msg = ChatMessage(session_id=session_id, role="assistant", content=content)
        db.add(msg)
        db.commit()
        db.refresh(msg)
        return msg.id
    finally:
        db.close()


@pytest.fixture(scope="module")
def client():
    import main

    importlib.reload(main)
    with TestClient(main.app) as c:
        yield c


@pytest.fixture(scope="module")
def auth(client):
    token = _register_and_login(client)
    message_id = _make_assistant_message(token, client)
    return {"token": token, "message_id": message_id}


def _get(client, auth, message_id, fmt):
    return client.get(
        f"/chat/messages/{message_id}/export",
        params={"format": fmt},
        headers=_auth_headers(auth["token"]),
    )


def test_export_unknown_format_422(client, auth):
    res = _get(client, auth, auth["message_id"], "exe")
    assert res.status_code == 422


def test_export_missing_message_404(client, auth):
    res = _get(client, auth, 99_999_999, "md")
    assert res.status_code == 404


def test_export_foreign_message_404(client, auth):
    other_token = _register_and_login(client)
    foreign_id = _make_assistant_message(other_token, client)
    # first user must not see the second user's message
    res = _get(client, auth, foreign_id, "md")
    assert res.status_code == 404


def test_export_md(client, auth):
    res = _get(client, auth, auth["message_id"], "md")
    assert res.status_code == 200, res.text
    assert res.headers["content-type"].startswith("text/markdown")
    assert "filename*=UTF-8''" in res.headers["content-disposition"]
    assert "| Направление |" in res.text


def test_export_txt_has_ascii_table(client, auth):
    res = _get(client, auth, auth["message_id"], "txt")
    assert res.status_code == 200
    body = res.content.decode("utf-8")
    assert "Направление" in body
    assert any(line.startswith("+") for line in body.splitlines())


def test_export_docx_opens_and_has_table(client, auth):
    from docx import Document

    res = _get(client, auth, auth["message_id"], "docx")
    assert res.status_code == 200
    doc = Document(io.BytesIO(res.content))
    assert len(doc.tables) == 1
    assert doc.tables[0].cell(0, 0).text.strip() == "Направление"


def test_export_pdf_or_clean_503(client, auth):
    import export_service

    res = _get(client, auth, auth["message_id"], "pdf")
    if export_service.pdf_available():
        assert res.status_code == 200
        assert res.content[:5] == b"%PDF-"
    else:
        assert res.status_code == 503
        assert "PDF" in res.json()["detail"]


def test_export_user_message_is_not_exportable(client, auth):
    """Only assistant answers are exportable."""
    from db import SessionLocal
    from models import ChatMessage, ChatSession

    db = SessionLocal()
    try:
        session_id = (
            db.query(ChatMessage.session_id)
            .filter(ChatMessage.id == auth["message_id"])
            .scalar()
        )
        user_msg = ChatMessage(session_id=session_id, role="user", content="привет")
        db.add(user_msg)
        db.commit()
        db.refresh(user_msg)
        user_msg_id = user_msg.id
    finally:
        db.close()
    res = _get(client, auth, user_msg_id, "md")
    assert res.status_code == 404
