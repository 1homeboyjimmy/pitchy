"""Export of chat answers and reports to PDF / DOCX / MD / TXT.

Pipeline: stored markdown → strip_llm_markup() → markdown→HTML → per-format
renderer. Shared by the message export endpoint (GET /chat/messages/{id}/export),
the chat-driven export flow («сохрани прошлый ответ в pdf») and the roadmap
PDF (roadmap_export.py).

PDF uses WeasyPrint + Jinja2 templates from export_templates/. WeasyPrint is
imported lazily: on bare Windows dev without GTK the import fails, so
pdf_available() gates the format and everything else keeps working.
DOCX builds real Word tables via python-docx; TXT renders aligned ASCII grids;
MD returns the cleaned markdown as-is.
"""

from __future__ import annotations

import io
import logging
import os
import re
import textwrap
from datetime import datetime
from pathlib import Path
from urllib.parse import quote, unquote, urlparse

from chat_attachments import sanitize_filename

logger = logging.getLogger("app")

EXPORT_FORMATS = ("pdf", "docx", "md", "txt")
EXPORT_MIMES = {
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "md": "text/markdown; charset=utf-8",
    "txt": "text/plain; charset=utf-8",
}
# Guards WeasyPrint/docx from pathological inputs; endpoint maps overflow to 413.
MAX_EXPORT_CHARS = 300_000

_TEMPLATES_DIR = Path(__file__).resolve().parent / "export_templates"
_PROJECT_DIR = Path(__file__).resolve().parent

# Marker persisted inside assistant message content so file cards survive
# reloads (same idea as the FILE markers in chat_attachments.py — keep the
# format in sync with parseExports in frontend/lib/utils.ts). Self-closing:
# message_id is present only when the marker points at ANOTHER message
# (export of a previous answer); without it the frontend uses the id of the
# message that carries the marker.
EXPORT_MARKER_RE = re.compile(
    r'<<<EXPORT format="([a-z]+)"(?: message_id="(\d+)")? name="([^"]*)">>>'
)


class ExportUnavailable(RuntimeError):
    """Requested format cannot be produced in this environment (PDF w/o GTK)."""


# ---------------------------------------------------------------------------
# Cleanup / markdown
# ---------------------------------------------------------------------------

_THOUGHT_TAGS = r"(?:think|thought|tool_call|tool_thought|think_process)"


def strip_llm_markup(content: str) -> str:
    """Server-side mirror of stripThoughts (frontend/lib/utils.ts) plus marker
    cleanup. Applied to everything that leaves the system as a document; also
    neutralizes active HTML since the result is embedded into PDF HTML."""
    if not content:
        return ""
    s = content
    s = re.sub(rf"<({_THOUGHT_TAGS})>[\s\S]*?</\1>", "", s, flags=re.I)
    s = re.sub(r"ǏǏǏ[\s\S]*?ǏǏǏ", "", s)
    s = re.sub(r"\[STATUS:[\s\S]*?\]", "", s)
    # Unclosed thought/status blocks (message cut mid-stream)
    s = re.sub(r"<(?:think|thought|think_process)>[\s\S]*$", "", s, flags=re.I)
    s = re.sub(r"ǏǏǏ[\s\S]*$", "", s)
    s = re.sub(r"\[STATUS:[\s\S]*$", "", s)
    # Active HTML: python-markdown passes raw HTML through, WeasyPrint renders it
    s = re.sub(r"<style[^>]*>[\s\S]*?</style>", "", s, flags=re.I)
    s = re.sub(r"<script[^>]*>[\s\S]*?</script>", "", s, flags=re.I)
    s = re.sub(r"<style[^>]*>[\s\S]*$", "", s, flags=re.I)
    s = re.sub(r"<script[^>]*>[\s\S]*$", "", s, flags=re.I)
    s = re.sub(r"<(/?)(iframe|object|embed|form|link|meta|base)\b", r"&lt;\1\2", s, flags=re.I)
    # Attachment / export markers embedded into stored messages
    s = re.sub(r'<<<FILE name="[^"]*" kind="[^"]*">>>\n?[\s\S]*?<<<END FILE>>>', "", s)
    s = EXPORT_MARKER_RE.sub("", s)
    return s.strip()


def markdown_to_html(md_text: str) -> str:
    """GFM-ish markdown → HTML. Mirrors the frontend renderer (react-markdown +
    remark-gfm): pipe tables, fenced code, sane list handling."""
    import markdown as _markdown

    return _markdown.markdown(
        md_text,
        extensions=["tables", "fenced_code", "sane_lists"],
    )


# ---------------------------------------------------------------------------
# PDF (WeasyPrint, lazy)
# ---------------------------------------------------------------------------

_pdf_state: dict[str, bool] = {"checked": False, "available": False}
_jinja_env_cache = None


def pdf_available() -> bool:
    """Cached probe: WeasyPrint imports only where its native libs exist
    (Docker image with pango; not bare Windows dev)."""
    if not _pdf_state["checked"]:
        try:
            import weasyprint  # noqa: F401

            _pdf_state["available"] = True
        except Exception as e:  # ImportError or OSError (missing gobject DLL)
            logger.warning(
                f"WeasyPrint unavailable, PDF export disabled: {type(e).__name__}: {e}"
            )
            _pdf_state["available"] = False
        _pdf_state["checked"] = True
    return _pdf_state["available"]


def _jinja_env():
    global _jinja_env_cache
    if _jinja_env_cache is None:
        from jinja2 import Environment, FileSystemLoader

        _jinja_env_cache = Environment(
            loader=FileSystemLoader(str(_TEMPLATES_DIR)),
            autoescape=True,
        )
    return _jinja_env_cache


def _safe_url_fetcher(url: str, timeout: int = 10, ssl_context=None):
    """Model output can smuggle <img src="http://…"> into the HTML — block all
    remote fetches (SSRF guard). Only project-local file:// (fonts, css) pass."""
    if url.startswith("file://"):
        from urllib.request import url2pathname

        path = Path(url2pathname(unquote(urlparse(url).path))).resolve()
        if path == _PROJECT_DIR or _PROJECT_DIR in path.parents:
            from weasyprint import default_url_fetcher

            return default_url_fetcher(url)
    raise ValueError(f"Blocked external resource in PDF export: {url[:120]}")


def html_to_pdf(html: str) -> bytes:
    """Full HTML document → PDF bytes. Raises ExportUnavailable without GTK."""
    if not pdf_available():
        raise ExportUnavailable("PDF export is not available in this environment")
    from weasyprint import HTML

    return HTML(
        string=html,
        base_url=str(_TEMPLATES_DIR) + os.sep,
        url_fetcher=_safe_url_fetcher,
    ).write_pdf()


def render_pdf_template(template_name: str, context: dict) -> bytes:
    if not pdf_available():
        raise ExportUnavailable("PDF export is not available in this environment")
    html = _jinja_env().get_template(template_name).render(**context)
    return html_to_pdf(html)


# ---------------------------------------------------------------------------
# DOCX (python-docx over the HTML tree)
# ---------------------------------------------------------------------------

def _docx_qn(tag):
    from docx.oxml.ns import qn

    return qn(tag)


def _docx_shade_cell(cell, fill: str) -> None:
    from docx.oxml import OxmlElement

    shd = OxmlElement("w:shd")
    shd.set(_docx_qn("w:val"), "clear")
    shd.set(_docx_qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def _docx_add_hyperlink(paragraph, url: str, text: str) -> None:
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.oxml import OxmlElement

    if not url or not url.lower().startswith(("http://", "https://", "mailto:")):
        paragraph.add_run(text)
        return
    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(_docx_qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(_docx_qn("w:val"), "1155CC")
    rpr.append(color)
    underline = OxmlElement("w:u")
    underline.set(_docx_qn("w:val"), "single")
    rpr.append(underline)
    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(_docx_qn("xml:space"), "preserve")
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _docx_inline(paragraph, node, bold=False, italic=False, code=False, skip=()):
    """Recursive inline content → runs. `skip` lets list items exclude their
    nested <ul>/<ol> (rendered separately as further list paragraphs)."""
    from bs4 import NavigableString
    from docx.shared import Pt

    for child in node.children:
        if isinstance(child, NavigableString):
            text = re.sub(r"\s+", " ", str(child))
            if not text.strip():
                continue
            run = paragraph.add_run(text)
            run.bold = bold
            run.italic = italic
            if code:
                run.font.name = "Consolas"
                run.font.size = Pt(9.5)
            continue
        name = getattr(child, "name", None)
        if name in skip:
            continue
        if name in ("strong", "b"):
            _docx_inline(paragraph, child, True, italic, code, skip)
        elif name in ("em", "i"):
            _docx_inline(paragraph, child, bold, True, code, skip)
        elif name == "code":
            _docx_inline(paragraph, child, bold, italic, True, skip)
        elif name == "a":
            _docx_add_hyperlink(paragraph, child.get("href", ""), child.get_text(" ", strip=True))
        elif name == "br":
            paragraph.add_run().add_break()
        else:
            _docx_inline(paragraph, child, bold, italic, code, skip)


def _docx_style(doc, paragraph, style_name: str) -> None:
    try:
        paragraph.style = doc.styles[style_name]
    except KeyError:
        pass


def _docx_list(doc, list_el, level: int = 0) -> None:
    base = "List Bullet" if list_el.name == "ul" else "List Number"
    style = base if level == 0 else f"{base} {min(level + 1, 3)}"
    for li in list_el.find_all("li", recursive=False):
        p = doc.add_paragraph()
        _docx_style(doc, p, style)
        _docx_inline(p, li, skip=("ul", "ol"))
        for sub in li.find_all(["ul", "ol"], recursive=False):
            _docx_list(doc, sub, level + 1)


def _docx_table(doc, table_el) -> None:
    rows = table_el.find_all("tr")
    if not rows:
        return
    n_cols = max(len(tr.find_all(["th", "td"])) for tr in rows)
    if n_cols == 0:
        return
    table = doc.add_table(rows=len(rows), cols=n_cols)
    try:
        table.style = doc.styles["Table Grid"]
    except KeyError:
        pass
    for i, tr in enumerate(rows):
        cells = tr.find_all(["th", "td"])
        header_row = bool(cells) and all(c.name == "th" for c in cells)
        for j in range(n_cols):
            cell = table.cell(i, j)
            if j >= len(cells):
                continue
            cell.text = ""
            _docx_inline(cell.paragraphs[0], cells[j], bold=header_row)
            if header_row:
                _docx_shade_cell(cell, "EEEEF2")
        if i == 0 and header_row:
            # Repeat the header row when the table spans pages
            from docx.oxml import OxmlElement

            tr_pr = table.rows[0]._tr.get_or_add_trPr()
            tbl_header = OxmlElement("w:tblHeader")
            tbl_header.set(_docx_qn("w:val"), "true")
            tr_pr.append(tbl_header)
    doc.add_paragraph()


def _docx_blocks(doc, parent) -> None:
    from bs4 import NavigableString

    for el in parent.children:
        if isinstance(el, NavigableString):
            text = str(el).strip()
            if text:
                doc.add_paragraph(text)
            continue
        name = el.name
        if name in ("h1", "h2", "h3", "h4", "h5", "h6"):
            heading = doc.add_heading("", level=min(int(name[1]), 4))
            _docx_inline(heading, el)
        elif name == "p":
            p = doc.add_paragraph()
            _docx_inline(p, el)
        elif name in ("ul", "ol"):
            _docx_list(doc, el)
        elif name == "table":
            _docx_table(doc, el)
        elif name == "pre":
            for line in el.get_text().rstrip("\n").split("\n"):
                p = doc.add_paragraph()
                run = p.add_run(line or " ")
                run.font.name = "Consolas"
                from docx.shared import Pt

                run.font.size = Pt(9)
        elif name == "blockquote":
            for sub in el.find_all(["p"], recursive=False) or [el]:
                p = doc.add_paragraph()
                _docx_style(doc, p, "Quote")
                _docx_inline(p, sub, italic=True)
        elif name == "hr":
            doc.add_paragraph()
        else:
            _docx_blocks(doc, el)


def html_to_docx(html: str, title: str = "", subtitle: str = "", sources=None) -> bytes:
    from bs4 import BeautifulSoup
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    if title:
        doc.add_heading(title, level=0)
    if subtitle:
        p = doc.add_paragraph()
        run = p.add_run(subtitle)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x90)
    soup = BeautifulSoup(html, "html.parser")
    _docx_blocks(doc, soup)
    if sources:
        doc.add_heading("Источники", level=2)
        for i, src in enumerate(sources, 1):
            p = doc.add_paragraph(f"{i}. ")
            title_text = (src.get("title") or src.get("url") or "").strip()
            _docx_add_hyperlink(p, src.get("url") or "", title_text or f"Источник {i}")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# TXT (aligned plain text with ASCII table grids)
# ---------------------------------------------------------------------------

_TXT_WIDTH = 100
_TXT_CELL_MAX = 40


def _txt_inline(el) -> str:
    """Inline text with useful degradation: links keep their URL."""
    from bs4 import NavigableString

    parts: list[str] = []
    for child in el.children:
        if isinstance(child, NavigableString):
            parts.append(str(child))
        elif child.name == "a":
            text = child.get_text(" ", strip=True)
            href = (child.get("href") or "").strip()
            parts.append(f"{text} ({href})" if href and href != text else text)
        elif child.name == "br":
            parts.append("\n")
        elif child.name in ("ul", "ol"):
            continue  # nested lists are rendered as separate lines
        else:
            parts.append(_txt_inline(child))
    return re.sub(r"[ \t]+", " ", "".join(parts)).strip()


def _txt_table(table_el) -> list[str]:
    rows_els = table_el.find_all("tr")
    rows = [[_txt_inline(c) for c in tr.find_all(["th", "td"])] for tr in rows_els]
    rows = [r for r in rows if r]
    if not rows:
        return []
    n_cols = max(len(r) for r in rows)
    rows = [r + [""] * (n_cols - len(r)) for r in rows]
    widths = [
        max(1, min(_TXT_CELL_MAX, max(len(row[j]) for row in rows)))
        for j in range(n_cols)
    ]
    header = bool(rows_els[0].find_all("th"))
    sep = "+" + "+".join("-" * (w + 2) for w in widths) + "+"
    head_sep = "+" + "+".join("=" * (w + 2) for w in widths) + "+"
    lines = [sep]
    for i, row in enumerate(rows):
        wrapped = [textwrap.wrap(cell, widths[j]) or [""] for j, cell in enumerate(row)]
        height = max(len(w) for w in wrapped)
        for k in range(height):
            cells = [
                (wrapped[j][k] if k < len(wrapped[j]) else "").ljust(widths[j])
                for j in range(n_cols)
            ]
            lines.append("| " + " | ".join(cells) + " |")
        lines.append(head_sep if header and i == 0 else sep)
    return lines


def _txt_list(list_el, out: list[str], level: int = 0) -> None:
    ordered = list_el.name == "ol"
    for idx, li in enumerate(list_el.find_all("li", recursive=False), 1):
        marker = f"{idx}. " if ordered else "- "
        indent = "  " * level
        text = _txt_inline(li)
        if text:
            out.extend(
                textwrap.wrap(
                    text,
                    _TXT_WIDTH,
                    initial_indent=indent + marker,
                    subsequent_indent=indent + " " * len(marker),
                )
                or [indent + marker]
            )
        for sub in li.find_all(["ul", "ol"], recursive=False):
            _txt_list(sub, out, level + 1)


def _txt_blocks(parent, out: list[str]) -> None:
    from bs4 import NavigableString

    for el in parent.children:
        if isinstance(el, NavigableString):
            text = str(el).strip()
            if text:
                out.extend(textwrap.wrap(text, _TXT_WIDTH))
            continue
        name = el.name
        if name in ("h1", "h2"):
            text = _txt_inline(el)
            out.extend(["", text.upper() if name == "h1" else text, ("=" if name == "h1" else "-") * min(len(text), _TXT_WIDTH), ""])
        elif name in ("h3", "h4", "h5", "h6"):
            out.extend(["", _txt_inline(el), ""])
        elif name == "p":
            text = _txt_inline(el)
            if text:
                for chunk in text.split("\n"):
                    out.extend(textwrap.wrap(chunk, _TXT_WIDTH) or [""])
                out.append("")
        elif name in ("ul", "ol"):
            _txt_list(el, out)
            out.append("")
        elif name == "table":
            out.extend(_txt_table(el))
            out.append("")
        elif name == "pre":
            out.extend("    " + line for line in el.get_text().rstrip("\n").split("\n"))
            out.append("")
        elif name == "blockquote":
            inner: list[str] = []
            _txt_blocks(el, inner)
            out.extend(("> " + line).rstrip() for line in inner)
            out.append("")
        elif name == "hr":
            out.extend(["-" * 40, ""])
        else:
            _txt_blocks(el, out)


def html_to_txt(html: str) -> str:
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(html, "html.parser")
    out: list[str] = []
    _txt_blocks(soup, out)
    text = "\n".join(line.rstrip() for line in out)
    return re.sub(r"\n{3,}", "\n\n", text).strip() + "\n"


# ---------------------------------------------------------------------------
# Filenames / Content-Disposition
# ---------------------------------------------------------------------------

_TRANSLIT = {
    "а": "a", "б": "b", "в": "v", "г": "g", "д": "d", "е": "e", "ё": "e",
    "ж": "zh", "з": "z", "и": "i", "й": "y", "к": "k", "л": "l", "м": "m",
    "н": "n", "о": "o", "п": "p", "р": "r", "с": "s", "т": "t", "у": "u",
    "ф": "f", "х": "h", "ц": "ts", "ч": "ch", "ш": "sh", "щ": "sch",
    "ъ": "", "ы": "y", "ь": "", "э": "e", "ю": "yu", "я": "ya",
}


def transliterate(text: str) -> str:
    out = []
    for ch in text:
        low = ch.lower()
        if low in _TRANSLIT:
            mapped = _TRANSLIT[low]
            out.append(mapped.capitalize() if ch.isupper() and mapped else mapped)
        else:
            out.append(ch)
    return "".join(out)


def suggest_filename(title: str | None, fmt: str) -> str:
    """Human filename (Cyrillic ok) from a session title / first heading."""
    base = re.sub(r"[#*_`\[\]<>|\"/\\:?]", "", title or "").strip()
    base = re.sub(r"\s+", " ", base)[:60].strip(" .-—–") or "Ответ Pitchy"
    date = datetime.now().strftime("%d.%m.%Y")
    return f"{base} — {date}.{fmt}"


def build_content_disposition(filename: str) -> str:
    """RFC 5987: ASCII fallback + UTF-8 filename* so Cyrillic names survive."""
    ascii_name = transliterate(filename)
    ascii_name = re.sub(r"[^A-Za-z0-9._ -]+", "", ascii_name)
    ascii_name = re.sub(r"\s+", " ", ascii_name).strip() or "export"
    return (
        f'attachment; filename="{ascii_name}"; '
        f"filename*=UTF-8''{quote(filename, safe='')}"
    )


# ---------------------------------------------------------------------------
# Export markers (chat-driven flow)
# ---------------------------------------------------------------------------

def build_export_marker(fmt: str, name: str, message_id: int | None = None) -> str:
    safe_name = sanitize_filename(name)
    if message_id is not None:
        return f'<<<EXPORT format="{fmt}" message_id="{int(message_id)}" name="{safe_name}">>>'
    return f'<<<EXPORT format="{fmt}" name="{safe_name}">>>'


# ---------------------------------------------------------------------------
# Chat-request detection (regex prefilter; the SLM precision pass lives in
# the chat pipeline and only runs when this returns True)
# ---------------------------------------------------------------------------

_EXPORT_FMT_RE = re.compile(
    r"(?:\bpdf\b|пдф|\bdocx?\b|ворд\w*|\bword\b|\btxt\b|тхт|текстов\w*|"
    r"\bmd\b|\bmarkdown\b|маркдаун\w*|файл\w*)",
    re.IGNORECASE,
)
_EXPORT_VERB_RE = re.compile(
    r"(?:сохран\w*|экспорт\w*|скач\w*|собер\w*|собра\w*|выгруз\w*|сдела\w*|"
    r"оформ\w*|сгенерир\w*|подготов\w*|запиш\w*|конверт\w*|"
    r"\bdownload\b|\bexport\b|\bsave\b|\bconvert\b)",
    re.IGNORECASE,
)


def detect_export_request(text: str) -> bool:
    """Cheap prefilter: a format word AND an action verb somewhere in the
    message. False positives are expected («сделай таблицу в markdown») —
    the SLM classifier decides; false negatives just mean no export offer."""
    if not text or len(text) > 2000:
        return False
    return bool(_EXPORT_FMT_RE.search(text) and _EXPORT_VERB_RE.search(text))


_EXPORT_INTENT_SYSTEM = """Ты определяешь, просит ли пользователь собрать ОТВЕТ ЧАТА В ФАЙЛ (pdf/docx/md/txt).
is_export=true только когда пользователь хочет получить файл с ответом ассистента.

target:
- "previous" — сохранить в файл уже данный (прошлый/последний/этот) ответ;
- "current" — в сообщении есть вопрос/задание, и его ответ надо сразу собрать в файл. Просьбы «следующий ответ собери в …» тоже считай current.

formats: перечисли запрошенные форматы строками из набора pdf, docx, md, txt.
Синонимы: ворд/word/вордовский → docx; текстовый файл/тхт → txt; маркдаун/markdown → md; пдф → pdf. Если формат не назван («сохрани в файл») — верни пустой список.

Примеры:
- «сохрани прошлый ответ в pdf» → is_export=true, formats=["pdf"], target="previous"
- «собери это в ворд и в txt» → is_export=true, formats=["docx","txt"], target="previous"
- «расскажи про юнит-экономику и оформи ответ в pdf» → is_export=true, formats=["pdf"], target="current"
- «сделай таблицу конкурентов в markdown» → is_export=false (просят формат текста ответа, не файл)
- «как экспортировать базу клиентов в excel?» → is_export=false (вопрос по теме, а не просьба собрать файл)
- «скачал pdf с презентацией, посмотри» → is_export=false"""


async def classify_export_intent(query: str, timeout: float = 3.5):
    """SLM precision pass (instructor + Qwen, same stack as dispatch_intent in
    llm_client.py) after detect_export_request(). Returns a normalized
    ExportIntent or None — None always degrades to the regular chat flow."""
    import asyncio

    from llm_client import get_instructor_client
    from schemas.llm import ExportIntent

    try:
        client = get_instructor_client("routerai")
        model = os.getenv("DISPATCHER_MODEL", "qwen/qwen-2.5-7b-instruct")
        res = await asyncio.wait_for(
            client.chat.completions.create(
                model=model,
                response_model=ExportIntent,
                messages=[
                    {"role": "system", "content": _EXPORT_INTENT_SYSTEM},
                    {"role": "user", "content": query[:2000]},
                ],
                max_retries=1,
            ),
            timeout=timeout,
        )
    except Exception as e:
        logger.warning(f"Export intent classification failed: {type(e).__name__}: {e}")
        return None
    if not res.is_export:
        return None
    formats = [f.lower().strip() for f in (res.formats or [])]
    formats = [f for f in dict.fromkeys(formats) if f in EXPORT_FORMATS] or ["pdf"]
    res.formats = formats
    res.target = "current" if res.target == "current" else "previous"
    return res


# ---------------------------------------------------------------------------
# Facade
# ---------------------------------------------------------------------------

def _sources_to_md(sources) -> str:
    lines = ["", "## Источники", ""]
    for i, src in enumerate(sources, 1):
        title = (src.get("title") or src.get("url") or f"Источник {i}").strip()
        url = (src.get("url") or "").strip()
        lines.append(f"{i}. [{title}]({url})" if url else f"{i}. {title}")
    return "\n".join(lines)


def render_message_export(
    content: str,
    fmt: str,
    title: str = "Ответ Pitchy",
    sources: list[dict] | None = None,
) -> tuple[bytes, str, str]:
    """Chat answer (markdown) → document bytes. Returns (data, mime, ext).

    Raises ValueError for unknown formats and ExportUnavailable when PDF is
    requested without WeasyPrint runtime libs.
    """
    if fmt not in EXPORT_FORMATS:
        raise ValueError(f"Unsupported export format: {fmt}")
    text = strip_llm_markup(content)[:MAX_EXPORT_CHARS]
    clean_sources = [
        {"title": str(s.get("title") or ""), "url": str(s.get("url") or "")}
        for s in (sources or [])
        if isinstance(s, dict) and (s.get("title") or s.get("url"))
    ]
    date_line = datetime.now().strftime("%d.%m.%Y")

    if fmt == "md":
        body = text
        if clean_sources:
            body += "\n" + _sources_to_md(clean_sources)
        return body.encode("utf-8"), EXPORT_MIMES[fmt], fmt

    html = markdown_to_html(text)

    if fmt == "txt":
        header = [title.upper(), f"{date_line} · Сгенерировано в Pitchy · pitchy.pro", "=" * _TXT_WIDTH, ""]
        body = html_to_txt(html)
        if clean_sources:
            src_lines = ["", "ИСТОЧНИКИ", "-" * 9]
            src_lines += [
                f"{i}. {(s['title'] or s['url'])} — {s['url']}" if s["title"] and s["url"] else f"{i}. {s['title'] or s['url']}"
                for i, s in enumerate(clean_sources, 1)
            ]
            body += "\n".join(src_lines) + "\n"
        return ("\n".join(header) + body).encode("utf-8"), EXPORT_MIMES[fmt], fmt

    if fmt == "docx":
        subtitle = f"{date_line} · Сгенерировано в Pitchy · pitchy.pro"
        data = html_to_docx(html, title=title, subtitle=subtitle, sources=clean_sources)
        return data, EXPORT_MIMES[fmt], fmt

    # pdf
    data = render_pdf_template(
        "message_pdf.html",
        {
            "title": title,
            "date": date_line,
            "body_html": html,
            "sources": clean_sources,
        },
    )
    return data, EXPORT_MIMES[fmt], fmt
