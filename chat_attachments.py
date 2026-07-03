"""Chat file/photo attachments: validation, text extraction, image processing.

Files are processed once at upload time (POST /chat/uploads): documents are
reduced to plain text, images are described by a vision model. The extracted
text is returned to the client, which sends it back together with the next
chat message; the message endpoint embeds it into the stored user message
between FILE markers so history reloads keep rendering the attachment chips.
"""

import io
import logging
import os
import re

import httpx

logger = logging.getLogger("app")

# Upload limits. Documents are text-extracted so the size mostly guards RAM;
# images are base64-encoded into the vision request, so keep them smaller.
MAX_UPLOAD_BYTES = 15 * 1024 * 1024
MAX_IMAGE_BYTES = 10 * 1024 * 1024
# Per-file cap on extracted text. GLM context fits far more, but the block is
# embedded into the stored message and round-trips through the client.
MAX_TEXT_CHARS = 15_000
MAX_FILES_PER_MESSAGE = 5
MAX_PDF_PAGES = 60

IMAGE_MIMES = {
    "image/png": "png",
    "image/jpeg": "jpg",
    "image/webp": "webp",
    "image/gif": "gif",
}

_EXT_KINDS = {
    "png": "image", "jpg": "image", "jpeg": "image", "webp": "image", "gif": "image",
    "pdf": "pdf",
    "docx": "docx",
    "txt": "text", "md": "text", "markdown": "text",
    "csv": "csv", "tsv": "csv",
    "json": "text",
}

# Markers used to embed extracted file text into the stored message content.
# The frontend parses these to render attachment chips instead of raw text
# (see parseAttachments in frontend/lib/utils.ts) — keep the formats in sync.
_FILE_MARKER_START = '<<<FILE name="{name}" kind="{kind}">>>'
_FILE_MARKER_END = "<<<END FILE>>>"


def sanitize_filename(name: str) -> str:
    """Strip path components and characters that would break the FILE marker."""
    name = (name or "file").replace("\\", "/").split("/")[-1]
    name = re.sub(r'["<>\r\n\t]', "", name).strip() or "file"
    return name[:120]


def detect_attachment_kind(filename: str, content_type: str | None) -> str | None:
    """Returns attachment kind ("image", "pdf", "docx", "text", "csv") or None."""
    if content_type and content_type.lower().split(";")[0] in IMAGE_MIMES:
        return "image"
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    return _EXT_KINDS.get(ext)


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "cp1251"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    if reader.is_encrypted:
        try:
            reader.decrypt("")
        except Exception:
            raise ValueError("PDF защищён паролем — снимите защиту и загрузите снова.")
    pages = []
    for page in reader.pages[:MAX_PDF_PAGES]:
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            pages.append("")
    text = "\n\n".join(p for p in pages if p.strip()).strip()
    if not text:
        raise ValueError(
            "Не удалось извлечь текст из PDF — похоже, это скан без текстового слоя. "
            "Попробуйте загрузить страницы как изображения."
        )
    return text


def _extract_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    text = "\n".join(parts).strip()
    if not text:
        raise ValueError("Документ не содержит текста.")
    return text


def extract_document_text(filename: str, kind: str, data: bytes) -> str:
    """Extracts plain text from a non-image upload. Raises ValueError with a
    user-facing (Russian) message when the file can't be processed."""
    if kind == "pdf":
        try:
            return _extract_pdf(data)
        except ValueError:
            raise
        except Exception as e:
            logger.error(f"PDF extraction failed for {filename}: {e}")
            raise ValueError("Не удалось прочитать PDF-файл. Убедитесь, что файл не повреждён.")
    if kind == "docx":
        try:
            return _extract_docx(data)
        except ValueError:
            raise
        except Exception as e:
            logger.error(f"DOCX extraction failed for {filename}: {e}")
            raise ValueError("Не удалось прочитать документ Word. Поддерживается только формат .docx.")
    if kind in ("text", "csv"):
        text = _decode_text(data).strip()
        if not text:
            raise ValueError("Файл пуст.")
        return text
    raise ValueError("Неподдерживаемый тип файла.")


_VISION_SYSTEM_PROMPT = (
    "Ты обрабатываешь изображения для ИИ-аналитика стартапов. Отвечай строго на русском. "
    "Опиши изображение подробно и структурированно: 1) что изображено; "
    "2) весь текст с изображения дословно; 3) все цифры, метрики, подписи осей и значения "
    "графиков или таблиц. Если это скриншот документа, презентации или переписки — передай "
    "содержимое максимально точно, без домыслов."
)


async def describe_image(data: bytes, mime: str, filename: str = "") -> str | None:
    """Runs the image through a vision model and returns the extracted
    description/text, or None when no vision provider is reachable.

    Providers are OpenAI-compatible; tried in order Makura → RouterAI so an
    outage or a missing vision model on one side degrades gracefully.
    """
    import base64

    data_url = f"data:{mime};base64,{base64.b64encode(data).decode()}"
    user_content = [
        {"type": "text", "text": f"Обработай изображение «{filename}»." if filename else "Обработай изображение."},
        {"type": "image_url", "image_url": {"url": data_url}},
    ]
    messages = [
        {"role": "system", "content": _VISION_SYSTEM_PROMPT},
        {"role": "user", "content": user_content},
    ]

    attempts = [
        (
            "https://api.makura.ai/v1",
            os.getenv("MAKURA_API_KEY"),
            os.getenv("MAKURA_VISION_MODEL", "glm-4.5v"),
        ),
        (
            os.getenv("ROUTERAI_API_BASE", os.getenv("ROUTERAI_BASE_URL", "https://routerai.ru/api/v1")),
            os.getenv("ROUTERAI_API_KEY"),
            os.getenv("ROUTERAI_VISION_MODEL", "gpt-4o-mini"),
        ),
    ]

    for base_url, api_key, model in attempts:
        if not api_key:
            continue
        try:
            async with httpx.AsyncClient(timeout=90) as client:
                resp = await client.post(
                    f"{base_url.rstrip('/')}/chat/completions",
                    headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
                    json={"model": model, "messages": messages, "temperature": 0.2, "max_tokens": 2000},
                )
            if resp.status_code != 200:
                logger.warning(f"Vision model {model} at {base_url} returned {resp.status_code}: {resp.text[:200]}")
                continue
            content = (resp.json().get("choices") or [{}])[0].get("message", {}).get("content") or ""
            if isinstance(content, list):
                content = "".join(str(c.get("text", "")) if isinstance(c, dict) else str(c) for c in content)
            if content.strip():
                logger.info(f"Vision extraction ok via {model} ({len(content)} chars)")
                return content.strip()
        except Exception as e:
            logger.warning(f"Vision model {model} at {base_url} failed: {type(e).__name__}: {e}")
    return None


def build_attachment_block(attachments) -> str:
    """Formats attachments (objects with .name/.kind/.text) into the marker
    block embedded into the stored user message content."""
    blocks = []
    for att in attachments[:MAX_FILES_PER_MESSAGE]:
        name = sanitize_filename(getattr(att, "name", "") or "file")
        kind = re.sub(r"[^a-z]", "", (getattr(att, "kind", "") or "file").lower())[:20] or "file"
        text = (getattr(att, "text", "") or "").strip()[:MAX_TEXT_CHARS]
        # A stray end-marker inside extracted text would truncate the block on parse.
        text = text.replace(_FILE_MARKER_END, "<END FILE>")
        blocks.append(f"{_FILE_MARKER_START.format(name=name, kind=kind)}\n{text}\n{_FILE_MARKER_END}")
    return "\n\n".join(blocks)
